import datetime
import os
import sys
import tkinter as tk
from tkinter import messagebox

from tkinter import ttk
import copy

import docx
import pandas as pd
import re


class FurnitureApp:
    def __init__(self, master):
        self.master = master
        self.lower_cabinets = self.load_cabinets("Szafki dolne.xlsx")
        self.upper_cabinets = self.load_cabinets("Szafki górne.xlsx")
        if self.lower_cabinets is None or self.upper_cabinets is None:
            sys.exit("Nie udało się załadować szafek.")
        self.cabinets = {**self.lower_cabinets, **self.upper_cabinets}
        self.cart = {}
        self.entry = None
        self.entry_kolor = None
        self.create_widgets()

    @staticmethod
    def load_cabinets(filename):
        if os.path.exists(filename):
            df = pd.read_excel(filename)
            df = df.fillna(0)
            df.to_excel(filename, index=False)
        else:
            messagebox.showwarning("Uwaga!", f"Nie znaleziono pliku {filename}")
            return None

        # group the DataFrame by the big table key column and convert each group to a dictionary
        grouped = df.groupby("Nazwa").apply(
            lambda x: x[
                x.columns.intersection(
                    ["partName", "height", "width", "pieces", "wrapping", "comments"]
                )
            ].to_dict("records"),
            include_groups=False,
        )
        cabinets = grouped.to_dict()
        return cabinets

    def create_widgets(self):
        parts_frame = ttk.Frame(self.master, padding=10)
        parts_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        parts_label = ttk.Label(parts_frame, text="Lista szafek", font=("Arial", 16))
        parts_label.pack(side=tk.TOP, pady=10)

        notebook = ttk.Notebook(parts_frame)
        notebook.pack(side=tk.LEFT, pady=10, expand=True)

        def split(a, n):
            k, m = divmod(len(a), n)
            return (
                a[i * k + min(i, m) : (i + 1) * k + min(i + 1, m)] for i in range(n)
            )

        column_1_low, column_2_low, column_3_low = split(
            list(self.lower_cabinets.keys()), 3
        )
        column_1_upp, column_2_upp, column_3_upp = split(
            list(self.upper_cabinets.keys()), 3
        )

        def create_buttons(column, column_frame):
            for cabinet in column:
                cabinet_frame = tk.Frame(column_frame, padx=10, pady=5)
                cabinet_frame.pack(side=tk.TOP, fill=tk.X, expand=True)
                cabinet_label = tk.Label(
                    cabinet_frame, text=cabinet, font=("Arial", 12)
                )
                cabinet_label.pack(side=tk.LEFT, padx=5)
                plus_button = tk.Button(
                    cabinet_frame,
                    text="+",
                    font=("Arial", 12),
                    width=2,
                    command=lambda selected=cabinet: self.add_to_cart(selected),
                )
                plus_button.pack(side=tk.RIGHT, padx=5)
                minus_button = tk.Button(
                    cabinet_frame,
                    text="-",
                    font=("Arial", 12),
                    width=2,
                    command=lambda selected=cabinet: self.remove_from_cart(selected),
                )
                minus_button.pack(side=tk.RIGHT, padx=5)

        parts_frame_low = tk.Frame(notebook, padx=10, pady=10)
        parts_frame_upp = tk.Frame(notebook, padx=10, pady=10)

        column_1_low_frame = tk.Frame(parts_frame_low, padx=10, pady=10)
        column_1_low_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        create_buttons(column_1_low, column_1_low_frame)

        column_2_low_frame = tk.Frame(parts_frame_low, padx=10, pady=10)
        column_2_low_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        create_buttons(column_2_low, column_2_low_frame)

        column_3_low_frame = tk.Frame(parts_frame_low, padx=10, pady=10)
        column_3_low_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        create_buttons(column_3_low, column_3_low_frame)

        column_1_upp_frame = tk.Frame(parts_frame_upp, padx=10, pady=10)
        column_1_upp_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        create_buttons(column_1_upp, column_1_upp_frame)

        column_2_upp_frame = tk.Frame(parts_frame_upp, padx=10, pady=10)
        column_2_upp_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        create_buttons(column_2_upp, column_2_upp_frame)

        column_3_upp_frame = tk.Frame(parts_frame_upp, padx=10, pady=10)
        column_3_upp_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        create_buttons(column_3_upp, column_3_upp_frame)

        notebook.add(parts_frame_low, text="Szafki dolne")
        notebook.add(parts_frame_upp, text="Szafki górne")

        cart_frame = tk.Frame(self.master, padx=10, pady=10)
        cart_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        default_text = "Nowe zamówienie " + datetime.date.today().strftime("%d-%m-%Y")
        default_kolor = "Kolor"

        cart_entry = tk.Entry(cart_frame, font=("Arial", 16), width=len(default_text))
        cart_entry.pack(side=tk.TOP, pady=10)
        cart_entry.insert(0, default_text)

        color_entry = tk.Entry(cart_frame, font=("Arial", 16), width=len(default_text))
        color_entry.pack(side=tk.TOP, pady=10)
        color_entry.insert(0, default_kolor)

        # create listbox for cart items and add it to the cart frame
        self.cart_listbox = tk.Listbox(cart_frame, font=("Arial", 12), height=10)
        self.cart_listbox.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        # create print button for cart and add it to the cart frame
        print_button = tk.Button(
            cart_frame,
            text="Zapisz",
            font=("Arial", 12),
            command=lambda cabinets=self.cabinets, entry_button=cart_entry, color_entry=color_entry: self.save_cart(
                cabinets, entry_button, color_entry
            ),
        )
        print_button.pack(side=tk.BOTTOM, pady=10, anchor=tk.SE)

    def add_to_cart(self, part):
        if part in self.cart:
            self.cart[part] += 1
        else:
            self.cart[part] = 1
        self.update_cart()

    def remove_from_cart(self, part):
        if part in self.cart:
            if self.cart[part] == 1:
                del self.cart[part]
            else:
                self.cart[part] -= 1
        self.update_cart()

    def update_cart(self):
        # clear cart listbox and add each cart item with count to it
        self.cart_listbox.delete(0, tk.END)
        for part, count in self.cart.items():
            self.cart_listbox.insert(tk.END, f"{part} ({count})")

    def save_cart(self, cabinets, entry_button, color_entry):
        # create a new Word document
        doc = docx.Document()
        user_entry = entry_button.get()
        user_entry_color = color_entry.get()
        # add a title to the document
        doc.add_heading(user_entry, level=0)
        paragraph = doc.add_paragraph(user_entry_color)
        run = paragraph.runs[-1]
        font = run.font
        font.name = "Arial"
        font.size = docx.shared.Pt(20)
        parts_list = []
        # add the contents of the cart to the document
        for cabinet, count in self.cart.items():
            partsTemp = copy.deepcopy(cabinets[cabinet])
            for partTemp in partsTemp:
                partTemp["pieces"] *= count
                parts_list.append(partTemp)
        parts_list_sorted = []

        for part in parts_list:
            found = False
            for existing_part in parts_list_sorted:
                if (
                    part["partName"] == existing_part["partName"]
                    and part["height"] == existing_part["height"]
                    and part["width"] == existing_part["width"]
                    and part["comments"] == existing_part["comments"]
                    and part["wrapping"] == existing_part["wrapping"]
                ):
                    existing_part["pieces"] += part["pieces"]
                    found = True
                    break
            if not found:
                parts_list_sorted.append(part)

        parts_list_sorted.sort(key=lambda x: x["partName"])

        def zero_to_str(wrapping):
            if wrapping == 0 or wrapping is None:
                wrapping = ""
            return wrapping

        for part in parts_list_sorted:
            doc.add_paragraph(
                f"{part['partName']}          {part['height']} X {part['width']}            {part['pieces']}szt       {zero_to_str(part['wrapping'])}        {zero_to_str(part['comments'])}"
            )

        # Remove special characters from user_entry
        user_entry = re.sub(r"\W+", "", user_entry)
        filename = user_entry + ".docx"
        # save the document to a file
        doc.save(filename)

        # open the saved file
        os.startfile(filename)


root = tk.Tk()

root.title("Formatki")
if os.path.exists("icon.ico"):
    root.iconbitmap("icon.ico")
app = FurnitureApp(root)
root.mainloop()

# Name of program rozpiski
# ikonka
# jesli word o takiej samej nazwie jest otwarty to zamkya
# moze jakies powiadomienia jakich plikow brakuje
